import auxil
import gen
import change_case
import consts

import os
import shutil
import subprocess as sb
from os.path import abspath
from random import randint, choice
import re
import copy

from docx import Document
from docx.shared import Mm, Pt
import json
from pdf2jpg import pdf2jpg
from PyPDF2 import PdfReader
from loguru import logger


ann_id = 0
# Добавляет ответственных и дедлайн в задачу, если их нет
def extend_instruction(instruction, samples_dir):
	for i in range(len(instruction)):
		instr = instruction[i]
		task_responsibles_people = instr["task_responsibles_people"]
		task_responsibles_groups = instr["task_responsibles_groups"]
		task_deadline = instr["task_deadline"]

		if (not task_responsibles_people) and (not task_responsibles_groups):
			
			# Шанс 25%
			if randint(1, 4) == 1:

				# Ответственные в новом предложении или в новом абзаце
				sep_char = ' ' if randint(0, 1) == True else '\n'

				with open(f"{samples_dir}/task_control.txt") as cfile:
					ctrl_list = cfile.read().split('\n')
					ctrl_msg = choice(ctrl_list)

				with open(f"{samples_dir}/responsible.json") as rfile:
					resp_list = json.load(rfile)
					resp = choice(resp_list)

				if resp[-1] != "group":
					instruction[i]["task_responsibles_people"] = resp[1:]
				else:
					instruction[i]["task_responsibles_groups"] = resp[1:-1]

				resp_to_doc = change_case.create_responsible(ctrl_msg, resp[0])
				instruction[i]["task_text"] += sep_char + resp_to_doc + '.'

				if ("оставляю за собой" in resp_to_doc):
					instruction[i]["task_responsibles_people"] = "Автор приказа."

		if (not task_deadline):

			# Шанс 25%
			if randint(1, 4) == 1:

				# Дедлайн в новом предложении или в новом абзаце
				sep_char = ' ' if randint(0, 1) == True else '\n'
				deadline = auxil.generate_date(unixtime=True)

				with open(f"{samples_dir}/task_deadline.txt") as ddfile:
					dd_list = ddfile.read().split('\n')
					deadline_msg = choice(dd_list)
					
				instruction[i]["task_deadline"] = deadline
				instruction[i]["task_text"] += sep_char + deadline_msg
				instruction[i]["task_text"] += deadline[0] + '.'

	return instruction

def write_docx(header, name, intro, instruction, responsible, creator,
		date, out, count, logo, sign, seal):
	document = Document()

	style = document.styles['Normal']
	font = style.font
	p_format = style.paragraph_format
	font.name = consts.font_name
	font.size = consts.font_size
	p_format.line_spacing_rule = consts.line_spacing

	for section in document.sections:
		section.top_margin = consts.top_margin
		section.bottom_margin =consts.bottom_margin
		section.left_margin = consts.left_margin
		section.right_margin = consts.right_margin

	if logo:
		document.add_picture(logo, consts.logo_w, consts.logo_h)

	headerp = document.add_paragraph()
	headerp.alignment = 1
	headerp.add_run(header + '\n\n').bold = True
	
	namep = document.add_paragraph()
	namep.alignment = 1
	namep.add_run(name)

	introp = document.add_paragraph(intro)
	introp.alignment = 1
	#introp.paragraph_format.first_line_indent = consts.first_line_indent
	
	instruction = auxil.add_numbering(instruction)
	for i in instruction:
		tsk = document.add_paragraph(i)
		tsk.alignment = 1
		#tsk.paragraph_format.first_line_indent = consts.first_line_indent
	responsiblep = document.add_paragraph(responsible)
	responsiblep.alignment = 1
	#responsiblep.paragraph_format.first_line_indent = consts.first_line_indent

	datep = document.add_paragraph(creator+'\n')
	datep.add_run(date)
	datep.alignment = 2

	if sign:
		signp = document.add_paragraph()
		signp.alignment = 2
		signr = signp.add_run()
		signr.add_picture(sign, consts.sign_w, consts.sign_h)

	if seal:
		sealp = document.add_paragraph()
		sealp.alignment = 2
		sealr = sealp.add_run()
		sealr.add_picture(seal, consts.seal_w, consts.seal_h)

	path = f"{out}/docx/{count}.docx"
	document.save(path)
	logger.debug(path)

	return path

def write_json(instruction, responsible_arr, date, out, count):
	json_dict = {
		"Tasks": {}
	}

	for i in range(len(instruction)):
		instr = instruction[i]
		task_text = instr["task_text"][4:].strip()
		task_responsibles_people = instr["task_responsibles_people"]
		task_responsibles_groups = instr["task_responsibles_groups"]
		task_deadline = instr["task_deadline"]

		json_dict["Tasks"][f"Task{i+1}"] = {"task_text": task_text}
		json_dict["Tasks"][f"Task{i+1}"]["task_responsibles_people"] = task_responsibles_people
		json_dict["Tasks"][f"Task{i+1}"]["task_responsibles_groups"] = task_responsibles_groups
		json_dict["Tasks"][f"Task{i+1}"]["task_deadline"] = task_deadline


	json_dict["Tasks"]["Global_supervisor"] = responsible_arr
	json_dict["Tasks"]["Global_deadline"] = date

	with open(f"{out}/json/{count}.json", "w") as jsonf:
		json.dump(json_dict, jsonf, ensure_ascii=False, indent=4)
		logger.debug(f"Saved {out}/json/{count}.json")

	return f"{out}/json/{count}.json"

# Создание и добавление аннотаций в формате MS-COCO 
def create_COCO_json (data,out):
	json_dict = {
		    "images": [],
		    "annotations": [],
			"categories": [{"supercategory": "", "id": 1, "name": "text"}, {"supercategory": "", "id": 2, "name": "title"}, {"supercategory": "", "id": 3, "name": "logo"}, {"supercategory": "", "id": 4, "name": "sign"}, {"supercategory": "", "id": 5, "name": "seal"},{"supercategory": "", "id": 6, "name": "date"},{"supercategory": "", "id": 7, "name": "name"}]
		}
	with open(f"{out}/json/annotations.json", "w") as jsonf:
		json.dump(json_dict, jsonf, ensure_ascii=False, indent=4)
		logger.debug(f"Saved {out}/json/annotations.json")
	return f"{out}/json/annotations.json"

def write_to_COCO_json(json_path, pdf_path, data, is_image=False):
	with open(json_path, "r") as json_file:
		json_dict = json.load(json_file)
	
	
	reader = PdfReader(pdf_path)
	name = os.path.basename(pdf_path)
	logger.warning(name)
	
	for i in range(len(reader.pages)):
		page = reader.pages[i]
		logger.warning(f"{i}page")
		
		#images dict
		image_inf = {}
		imgid = str(i) + name.split('.')[0]
		image_inf['file_name'] = f"{i}_{name}.jpg"
		#pg_box = reader.pages[i].mediabox
		image_inf['width'] = 2550#pg_box.width
		image_inf['height'] = 3300#pg_box.height
		image_inf['id'] = imgid
		json_dict["images"].append(image_inf)
		
		
		
		#annotations dict ; classes = ['text', 'title', 'logo', 'sign', 'seal','date','name']
		global ann_id
		ann_inf = {}
		ann_inf['iscrowd'] = 0
		ann_inf['image_id'] = imgid
		
		if is_image:			
			# Координаты логотипа
			if i == 0:
				bbox,area = auxil.calculate_logo_coords()
				ann_inf['area'] = area
				ann_inf['bbox'] = bbox
				ann_inf['category_id'] = 3
				ann_inf['id'] = ann_id
				logger.debug(ann_inf)
				cp = copy.deepcopy(ann_inf)
				json_dict["annotations"].append(cp)
				ann_id += 1

			# Координаты подписи
			
			
			(tmx, tmy, im_count) = extract_tm(pdf_path, -1)

			if im_count >= 2: # Если на последней странице есть подпись и печать

				if (tmx == 0) or (tmy == 0): # Если на странице только подпись и печать
					sign_coords, area = auxil.calculate_sign_coords(tmx, tmy, new_page=True)
					
				else:
					sign_coords, area = auxil.calculate_sign_coords(tmx, tmy)
				right_page = len(reader.pages) - 1
				
				
			else: # Если на последней странице только печать
				(tmx, tmy, _) = extract_tm(pdf_path, -2)
				sign_coords, area = auxil.calculate_sign_coords(tmx, tmy)
				right_page = len(reader.pages) - 2
				
			if i == right_page:
				ann_inf['area'] = area
				ann_inf['bbox'] = sign_coords
				ann_inf['category_id'] = 4
				ann_inf['id'] = ann_id
				logger.debug(ann_inf)
				cp = copy.deepcopy(ann_inf)
				json_dict["annotations"].append(cp)
				ann_id += 1
					
				# Координаты печати
			if i == len(reader.pages) - 1:
				if im_count >= 2:
					seal_coords, area = auxil.calculate_seal_coords(sign_coords)

				else:
					seal_coords, area = auxil.calculate_seal_coords([], new_page=True)
				
				ann_inf['area'] = area
				ann_inf['bbox'] = seal_coords
				ann_inf['category_id'] = 5
				ann_inf['id'] = ann_id
				logger.debug(ann_inf)
				cp = copy.deepcopy(ann_inf)
				json_dict["annotations"].append(cp)
				ann_id += 1
				
		#text cord
		text_coords = auxil.calculate_text_coords(pdf_path, data, page)
		for j in range(len (text_coords)):
			if(text_coords[j] != ["page_break"]):
				if(j in {0,1}):
					ann_inf['category_id'] = 2
				elif j == 5:
					ann_inf['category_id'] = 7
				elif j == 6:
					ann_inf['category_id'] = 6
				else:
					ann_inf['category_id'] = 1
				
				for k in range(len(text_coords[j]) - 1):
					if(text_coords[j][k] != []):
						ann_inf['area'] = text_coords[j][k][1]
						ann_inf['bbox'] = text_coords[j][k][0]
						ann_inf['id'] = ann_id
						ann_id += 1
						logger.debug(ann_inf)
						cp = copy.deepcopy(ann_inf)
						json_dict["annotations"].append(cp)
				
		
	with open(json_path, "w") as jsonf:	
		json.dump(json_dict, jsonf, ensure_ascii=False, indent=4)
		

def write_pdf_linux(docx_path, out, count):
	out_path = f"{out}/pdf/{count}.pdf"
	out_path = abspath(out_path)
	docx_path = abspath(docx_path)

	cmd = ""
	cmd += f"abiword "
	cmd += f"-t pdf "
	cmd += f"-o {out_path} "
	cmd += docx_path
	sb.call(cmd, shell=True, stderr=sb.DEVNULL)
	logger.debug(f"Saved {out_path}")

	return out_path

#После преобразования из pdf в jpg все страницы складируются в одну папку.
def write_jpg(out, count):
	
	pdf = f"{out}/pdf/{count}.pdf"
	output = f"{out}/jpg"

	pdf2jpg.convert_pdf2jpg(pdf, output, pages="ALL")
	
	folder = os.path.join(output,f"{count}.pdf_dir")
	file_names = os.listdir(folder)
    
	for file_name in file_names:
		shutil.move(os.path.join(folder, file_name), output)
		
		if randint(1,5) == 1:
			auxil.add_noise(os.path.join(output, file_name), randint(70, 90))
			

		
		
	os.rmdir(folder)
	logger.debug(f"Saved {out}/jpg")

def extract_tm(pdf_path, page_num):
	reader = PdfReader(pdf_path)
	page = reader.pages[page_num]

	tmx, tmy = [], []
	def visitor_sign(text, cm, tm, fontDict, fontSize):
		tmx.append(tm[4])
		tmy.append(tm[5])

	text = page.extract_text(visitor_text=visitor_sign)

	# Координаты последнего блока с текстом
	# в PdfUnits
	tmx = tmx[-1]
	tmy = tmy[-1]

	return (tmx, tmy, len(page.images))

def write_coords(json_path, pdf_path, data, is_image=False):

	with open(json_path, "r") as json_file:
		json_dict = json.load(json_file)

	if is_image:

		# Координаты логотипа
		logo_coords = auxil.calculate_logo_coords()

		# Координаты подписи
		(tmx, tmy, im_count) = extract_tm(pdf_path, -1)

		if im_count >= 2: # Если на последней странице есть подпись и печать

			if (tmx == 0) or (tmy == 0): # Если на странице только подпись и печать
				sign_coords = auxil.calculate_sign_coords(tmx, tmy, new_page=True)
				
			else:
				sign_coords = auxil.calculate_sign_coords(tmx, tmy)

		else: # Если на последней странице только печать
			(tmx, tmy, _) = extract_tm(pdf_path, -2)
			sign_coords = auxil.calculate_sign_coords(tmx, tmy)

		# Координаты печати
		if im_count >= 2:
			seal_coords = auxil.calculate_seal_coords(sign_coords)

		else:
			seal_coords = auxil.calculate_seal_coords([], new_page=True)
		
		json_dict["Images"] = {}
		json_dict["Images"]["logo_coordinates"] = logo_coords
		json_dict["Images"]["signature_coordinates"] = sign_coords
		json_dict["Images"]["seal_coordinates"] = seal_coords

	text_coords = auxil.calculate_text_coords(pdf_path, data)
	json_dict["Text"] = {}
	json_dict["Text"]["header"] = text_coords[0]
	json_dict["Text"]["name"] = text_coords[1]
	json_dict["Text"]["intro"] = text_coords[2]
	json_dict["Text"]["tasks"] = text_coords[3]
	json_dict["Text"]["responsible"] = text_coords[4]
	json_dict["Text"]["creator"] = text_coords[5]
	json_dict["Text"]["date"] = text_coords[6]

	with open(json_path, "w") as jsonf:
		json.dump(json_dict, jsonf, ensure_ascii=False, indent=4)
